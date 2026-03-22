from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pathlib import Path

DOC_PATH = Path('/Users/celmargalindez/GABRIEL DOMINGO/pos_system/docs/Charlie-PC-System-Integration-and-Architecture-Panel.docx')
FIG_DIR = Path('/Users/celmargalindez/GABRIEL DOMINGO/pos_system/docs/figures')


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, line=1.5):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, before=0, after=8):
    p = doc.add_paragraph()
    format_paragraph(p, align=align, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    size_map = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    format_paragraph(p, align=align, before=12, after=8, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=size_map.get(level, 12), bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        format_paragraph(p)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        format_paragraph(p)
        run = p.add_run(item)
        set_run_font(run)


def add_table_title(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6, line=1.2)
    run = p.add_run(text)
    set_run_font(run, italic=True)


def shade_cell(cell, fill='F2F2F2'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        shade_cell(hdr_cells[i])
        for paragraph in hdr_cells[i].paragraphs:
            format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.2)
            for run in paragraph.runs:
                set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                format_paragraph(paragraph, line=1.2)
                for run in paragraph.runs:
                    set_run_font(run)
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_figure(doc, image_path, caption, width_inches):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=6, line=1.0)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    format_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.2)
    run = cap.add_run(caption)
    set_run_font(run, size=11)


def apply_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)


def build_document():
    doc = Document()
    apply_document_defaults(doc)

    # Title page
    for line in [
        'Republic of the Philippines',
        '[Name of School / University]',
        '[College / Department]'
    ]:
        add_text(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    for _ in range(3):
        add_text(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_heading(doc, 'CHARLIE PC WEB AND MOBILE SYSTEM INTEGRATION AND ARCHITECTURE DOCUMENTATION', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'A Formal System Integration and Architecture Report for the Charlie PC Retail Platform', align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        add_text(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, 'A Documentation Report Presented to the Panel', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, 'in Partial Fulfillment of the Requirements for', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, '[Course / Subject Title]', align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        add_text(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, 'Submitted by:', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, '[Name of Proponent 1]', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, '[Name of Proponent 2]', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, '[Name of Proponent 3]', align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        add_text(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, 'Adviser:', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, '[Name of Adviser]', align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        add_text(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(doc, 'March 2026', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)

    # Approval sheet
    add_heading(doc, 'APPROVAL SHEET', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'This documentation entitled “Charlie PC Web and Mobile System Integration and Architecture Documentation,” prepared and submitted by the proponents named above, is hereby presented to the panel as evidence of the system integration, design decisions, and architectural implementation of the Charlie PC retail platform.')
    add_text(doc, 'The undersigned certify that this report has been reviewed and recommended for acceptance in partial fulfillment of the academic requirements for the course or subject indicated by the institution.')
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    signatures = [
        ('[Name of Adviser]\nAdviser', '[Name of Panel Chair]\nPanel Chair'),
        ('[Name of Panel Member 1]\nPanel Member', '[Name of Panel Member 2]\nPanel Member'),
        ('[Name of Dean / Program Head]\nDean / Program Head', '')
    ]
    for r, pair in enumerate(signatures):
        for c, text in enumerate(pair):
            cell = sig_table.rows[r].cells[c]
            cell.text = '\n\n\n' + text
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
                for run in paragraph.runs:
                    set_run_font(run)
    add_text(doc, 'Date Approved: ________________________________', align=WD_ALIGN_PARAGRAPH.CENTER, before=18)
    add_page_break(doc)

    # Abstract
    add_heading(doc, 'ABSTRACT', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, 'Charlie PC is an integrated retail system developed for a computer parts and accessories shop that requires one consistent platform for customer ordering, cashier point-of-sale processing, and administrative monitoring. The implemented solution consists of a web application, a mobile application, and a shared Express-based backend connected to an SQLite database. The system centralizes authentication, inventory, checkout computation, receipt generation, reporting, and role-based access control so that both client applications follow the same operational rules.')
    add_text(doc, 'This report documents the architecture and integration strategy of the Charlie PC system. It explains how the web and mobile interfaces communicate with the backend through REST-style JSON endpoints, how JWT-based authentication supports customers, cashiers, and administrators, and how backend validation protects stock accuracy, pricing, payment rules, and receipt records. The report also presents the actual implemented modules, including realistic checkout flows, Philippine customer-data validation, delivery address selection, and responsive mobile layouts.')
    add_text(doc, 'Findings from the implemented system show that a layered client-server architecture is appropriate for a small retail environment that requires synchronized data across multiple interfaces. The chosen design provides a practical academic demonstration of web and mobile integration while remaining extensible for future improvements such as live payment-gateway integration, wider address datasets, and stronger deployment-scale infrastructure.')
    add_text(doc, 'Keywords: system integration, retail system, point-of-sale, mobile application, web application, REST API, architecture', bold=True, after=0)
    add_page_break(doc)

    # TOC and LOF
    add_heading(doc, 'TABLE OF CONTENTS', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for line in [
        'Title Page ............................................................................................................. i',
        'Approval Sheet ..................................................................................................... ii',
        'Abstract ............................................................................................................... iii',
        'List of Figures ..................................................................................................... iv',
        'Chapter I. Project Overview ................................................................................. 1',
        'Chapter II. System Architecture ............................................................................ 4',
        'Chapter III. System Integration and Design ........................................................... 8',
        'Chapter IV. Summary, Conclusion, and Recommendations .................................. 13',
        'Appendix A. Implementation Alignment .............................................................. 15',
        'References .......................................................................................................... 17'
    ]:
        add_text(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_heading(doc, 'LIST OF FIGURES', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for line in [
        'Figure 1. Charlie PC Web Customer Storefront .................................................... 6',
        'Figure 2. Charlie PC Web Admin Dashboard ....................................................... 7',
        'Figure 3. Charlie PC Mobile Sign-In Screen ........................................................ 12'
    ]:
        add_text(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_page_break(doc)

    # Chapter I
    add_heading(doc, 'Chapter I. Project Overview', level=1)
    add_heading(doc, 'I. Background of the Study', level=2)
    add_text(doc, 'Charlie PC is an integrated retail system designed for a computer parts and accessories shop that serves three main user groups: customers, cashiers, and administrators. The project was developed as a shared platform with one backend service and two client applications: a web application for customer ordering, cashier point-of-sale, and admin management, and a mobile application for customer shopping, cashier counter operations, and admin management.')
    add_text(doc, 'The practical need behind the system is common in small and medium retail stores. Many stores still separate customer ordering, cashier checkout, stock monitoring, and sales reporting into disconnected tools or manual records. This causes duplicated product encoding, inconsistent stock levels, and delayed reporting. Charlie PC addresses that problem by centralizing authentication, inventory, sales processing, and reporting in one backend while exposing the same business rules to both web and mobile clients.')
    add_text(doc, 'The system follows a layered client-server model and exposes REST-style JSON endpoints so that different clients can communicate with the same server logic in a consistent way, which is aligned with network-based architectural principles described by Fielding (2000).')
    add_heading(doc, 'II. Problem Statement', level=2)
    add_numbered(doc, [
        'Retail operations are harder to manage when customer ordering, cashier checkout, and inventory tracking are handled in separate tools.',
        'Product, price, and stock data become inconsistent when web and mobile interfaces do not share a single source of truth.',
        'Manual or client-side-only checkout logic can lead to invalid totals, stock mismatches, and weak auditability.',
        'Stores need role-based access so that customers, cashiers, and administrators see only the functions appropriate to their responsibilities.',
        'Small stores need a practical solution that supports web and mobile access without requiring a complex enterprise database setup.'
    ])
    add_heading(doc, 'III. Objectives', level=2)
    add_heading(doc, 'General Objective', level=3)
    add_text(doc, 'To design and implement an integrated Charlie PC system that connects web and mobile applications to a shared backend for authentication, product management, checkout, sales recording, and reporting.')
    add_heading(doc, 'Specific Objectives', level=3)
    add_numbered(doc, [
        'To provide a web application that supports customer browsing, online ordering, cashier POS operations, and admin product and sales management.',
        'To provide a mobile application that supports customer shopping, cashier counter operations, and admin dashboards, product management, and sales monitoring.',
        'To centralize authentication using token-based login so that both web and mobile clients use the same account and authorization rules.',
        'To centralize inventory and checkout rules so that sale totals, stock deductions, receipts, and validation are computed on the server.',
        'To maintain one consistent database for users, products, sales, and sale items.',
        'To support realistic checkout features such as receipts, delivery and pickup selection, PH customer-data validation, and multiple payment methods including Cash, Card, GCash, and Cash on Delivery.',
        'To keep the mobile application usable across narrow phones, tablets, and foldable-width layouts through responsive screen composition.'
    ])
    add_heading(doc, 'IV. Scope and Limitations', level=2)
    add_heading(doc, 'Scope', level=3)
    add_bullets(doc, [
        'Customer registration and login',
        'Role-based login for customer, cashier, and admin',
        'Product listing, search, and category browsing',
        'Add-to-cart and customer checkout on web and mobile',
        'Cashier POS checkout with quick cash tendering, payment confirmation, and receipt output',
        'Admin product management and sales monitoring on web and mobile',
        'Sales reporting through dashboard metrics',
        'PH phone, email, and delivery address validation with inline required-field feedback on web and mobile',
        'Receipt generation and inventory deduction',
        'Responsive mobile layouts that adapt for small phones, larger phones, tablets, and foldable-width screens',
        'Shared backend integration for both web and mobile clients'
    ])
    add_heading(doc, 'Limitations', level=3)
    add_bullets(doc, [
        'The backend uses SQLite, which is practical for a school project or small deployment, but is not intended for large multi-branch concurrency.',
        'Card and GCash flows are realistic UI and validation flows, but they are not yet connected to a live payment gateway or webhook-based payment confirmation.',
        'The Philippine address dataset is currently a curated starter dataset, not a complete nationwide PSGC-integrated source.',
        'The mobile cart is stateful during app use but does not yet implement advanced offline sync.',
        'The application includes validation and UI safeguards, but it does not yet use a full automated CI pipeline for every build target.',
        'The system currently runs as a centralized single backend service and is not yet deployed as distributed microservices.'
    ])
    add_page_break(doc)

    # Chapter II
    add_heading(doc, 'Chapter II. System Architecture', level=1)
    add_heading(doc, 'I. Architectural Style', level=2)
    add_text(doc, 'The Charlie PC system uses a hybrid client-server, layered three-tier, REST API-based architecture. The web and mobile applications act as the presentation layer. The Express backend provides the application and service layer, while SQLite and reference datasets act as the data layer. This architectural choice allows both clients to rely on one source of truth for authentication, products, checkout, receipts, and reporting.')
    add_table_title(doc, 'Table 1. Charlie PC Architecture Layers')
    add_table(doc, ['Layer', 'Main Components', 'Implemented Responsibility'], [
        ['Presentation Layer', 'React + Vite web client; Expo + React Native mobile client', 'Customer storefront, cashier counter flow, admin dashboard, role-based navigation, and form interaction'],
        ['Application / Service Layer', 'Express REST API, JWT authentication, role middleware, product service, sales service, reports service, reference-data service', 'Centralized request handling, validation, checkout computation, receipt generation, and reporting'],
        ['Data Layer', 'SQLite database, PH address reference data, automatic backups', 'Persistent storage for users, products, sales, sale items, and structured delivery metadata']
    ])
    add_heading(doc, 'II. Operational Flow', level=2)
    add_text(doc, 'The implemented system follows one common transaction flow even though users interact through different roles and interfaces. Customers browse products and may continue to cart and checkout only when authenticated as customers. Cashiers process over-the-counter transactions through the POS interface, while administrators monitor the system, manage products, and review sales records. In each case, the clients submit requests to the backend, the backend validates identity and business rules, and the resulting records are written to the shared database.')
    add_numbered(doc, [
        'The user opens Charlie PC using the web application or the mobile application.',
        'The user authenticates based on the appropriate role: customer, cashier, or administrator.',
        'The client collects the necessary product, payment, and fulfillment details.',
        'The backend validates the token, role, stock levels, payment rules, customer details, and checkout totals.',
        'If validation succeeds, the backend stores the sale, inserts line items, deducts stock, and generates a receipt number.',
        'If validation fails, the backend returns structured error feedback to the client.'
    ])
    add_heading(doc, 'III. Actual User Interface Figures', level=2)
    add_text(doc, 'The following figures present selected screenshots of the implemented Charlie PC system. These figures were captured from the current web application and the mobile application interface build used by the project.')
    add_figure(doc, FIG_DIR / 'web-customer-storefront.png', 'Figure 1. Charlie PC web customer storefront showing the public shopping interface and product discovery area.', 6.0)
    add_figure(doc, FIG_DIR / 'web-admin-dashboard.png', 'Figure 2. Charlie PC web admin dashboard showing operational metrics, revenue visibility, and administrative controls.', 6.0)
    add_figure(doc, FIG_DIR / 'mobile-login-screen.png', 'Figure 3. Charlie PC mobile sign-in screen showing the role-aware mobile entry point for customer, cashier, and admin users.', 3.0)
    add_heading(doc, 'IV. Data Model Summary', level=2)
    add_table_title(doc, 'Table 2. Core Operational Data Entities')
    add_table(doc, ['Entity', 'Selected Attributes', 'Purpose in the Implemented System'], [
        ['Users', 'ID, name, email, password, role', 'Stores role-based accounts for customers, cashiers, and administrators'],
        ['Products', 'ID, name, category, price, stock, barcode, image', 'Stores the inventory catalog used by web and mobile clients'],
        ['Sales', 'ID, total, subtotal, receipt number, payment method, fulfillment type, status, customer details', 'Stores receipt-level transaction details, payment metadata, fulfillment metadata, and audit status'],
        ['Sale Items', 'ID, sale ID, product ID, product name snapshot, quantity, price, line total', 'Stores the normalized line items associated with each sale and preserves receipt snapshots']
    ])
    add_page_break(doc)

    # Chapter III
    add_heading(doc, 'Chapter III. System Integration and Design', level=1)
    add_heading(doc, 'I. Integration Strategy', level=2)
    add_text(doc, 'The Charlie PC system integrates the web and mobile applications through a shared backend API and token-based authentication. Both clients consume the same backend endpoints, including /api/login, /api/register, /products, /sales, /reports/dashboard, and /reference/ph-addresses. This design allows the user interface to vary by platform while the system behavior remains consistent.')
    add_text(doc, 'After login, the backend returns a JWT containing the authenticated user ID and role. The web client stores the token in local storage, while the mobile client persists it through AsyncStorage. Protected requests send the token back in the Authorization header, allowing the backend to enforce access rules for customer checkout, cashier POS operations, and admin-level management tasks.')
    add_heading(doc, 'II. Shared Business Rules', level=2)
    add_bullets(doc, [
        'Customer online checkout is allowed only for authenticated customer accounts.',
        'POS checkout is allowed only for authenticated cashier or administrator accounts.',
        'Product creation is limited to authenticated admin or cashier users, while update and delete actions remain admin-only.',
        'Totals, discounts, tax, stock checks, receipt numbers, and sale items are computed server-side.',
        'Philippine email and phone validation are enforced before checkout data is accepted.',
        'Delivery checkout requires a structured address, and Cash on Delivery is limited to delivery orders.',
        'GCash requires a payment reference, while card payments preserve only masked last-four metadata.',
        'Voiding restocks items and marks the sale as voided instead of deleting the transaction record.'
    ])
    add_heading(doc, 'III. Validation and Reference Data Flow', level=2)
    add_text(doc, 'Validation is handled in both the client and the backend so that the interface can guide users while the server remains the final source of truth. The web and mobile checkout forms mark required fields, validate missing data, and restrict invalid combinations such as Cash on Delivery for pickup orders. The backend re-validates submitted values before accepting the transaction.')
    add_text(doc, 'Delivery addresses follow a structured Philippine address flow consisting of Region, Province, Municipality or City, Barangay, and Postal Code. The user then adds street, building, or landmark details before the final address is composed and stored with the sale.')
    add_heading(doc, 'IV. Database Integration', level=2)
    add_text(doc, 'SQLite was selected as the operational database because it is lightweight, zero-configuration, ACID-compliant, and practical for a capstone or small single-store deployment. The backend initializes and migrates the schema on startup, and the applications never connect to the database directly. Sales are created inside a transaction so that line-item insertion, stock deduction, and receipt creation happen atomically. The system also creates timestamped SQLite backups during startup to strengthen data protection for demonstration and development use.')
    add_heading(doc, 'V. Technology Stack', level=2)
    add_table_title(doc, 'Table 3. Major Technologies Used in Charlie PC')
    add_table(doc, ['System Part', 'Technology', 'Role in the Project'], [
        ['Web Application', 'React, Vite, Fetch API, custom Material-inspired styling', 'Browser-based storefront, cashier POS, and admin management interface'],
        ['Mobile Application', 'Expo SDK 54, React Native, React Navigation, React Native Paper, AsyncStorage', 'Cross-platform role-based mobile interface with responsive layouts and Material Design 3-style components'],
        ['Backend', 'Node.js, Express, jsonwebtoken, bcrypt, SQLite', 'Authentication, authorization, product CRUD, checkout processing, reporting, and reference-data handling']
    ])
    add_page_break(doc)

    # Chapter IV
    add_heading(doc, 'Chapter IV. Summary, Conclusion, and Recommendations', level=1)
    add_heading(doc, 'I. Summary', level=2)
    add_text(doc, 'The Charlie PC project successfully integrates a web application and a mobile application using one shared backend and one shared database. The system supports customer shopping, cashier POS operations, and admin management through a consistent role-based model. The architecture is best described as a layered client-server system with REST-style API communication. This design allows the project to reuse backend logic across both clients, keep critical business rules centralized, and maintain a consistent user experience across browser-based and mobile interfaces.')
    add_heading(doc, 'II. Conclusion', level=2)
    add_text(doc, 'Based on the implemented structure of the project, the chosen architecture is appropriate for the Charlie PC system requirements. Both clients use the same API, authentication is centralized using JWT, inventory and checkout rules are computed on the backend, and product, sales, and reporting data remain synchronized across platforms. Therefore, the project demonstrates a practical implementation of web-mobile integration for a retail environment and is suitable for academic presentation, prototype deployment, and small store operations.')
    add_heading(doc, 'III. Recommendations', level=2)
    add_numbered(doc, [
        'Integrate a real payment gateway for card and GCash processing.',
        'Replace the curated address dataset with a complete and authoritative Philippine administrative dataset.',
        'Add refresh-token or session-hardening support for stronger authentication lifecycle management.',
        'Add persistent cart synchronization across sessions and devices.',
        'Introduce automated CI validation that runs web build checks, mobile guard checks, and backend tests together.',
        'Migrate to a server-grade database if the project is expanded into a multi-branch deployment.',
        'Add barcode scanning, supplier modules, and purchase-order management for a more complete retail workflow.'
    ])
    add_page_break(doc)

    # Appendix and references
    add_heading(doc, 'Appendix A. Implementation Alignment', level=1)
    add_text(doc, 'The following files represent the major implementation points discussed in this report and demonstrate how the documented architecture maps to the actual Charlie PC codebase.')
    add_bullets(doc, [
        'frontend/src/App.jsx - web entry and role flow',
        'frontend/src/components/CustomerApp.jsx - web customer storefront and customer checkout',
        'frontend/src/components/POS.jsx - web cashier POS and payment flow',
        'frontend/src/components/Dashboard.jsx, frontend/src/components/ProductsPage.jsx, frontend/src/components/SalesPage.jsx - web admin interface',
        'frontend/src/constants/authRoles.js - shared web auth role definitions',
        'mobile/App.js - mobile app root',
        'mobile/src/navigation/AppNavigator.js - role-based mobile navigation',
        'mobile/src/screens/LoginScreen.js - mobile sign-in interface',
        'mobile/src/screens/CartScreen.js - mobile checkout and cashier payment flow',
        'mobile/src/screens/ProductsScreen.js - mobile shop screen',
        'mobile/src/screens/AdminDashboardScreen.js, mobile/src/screens/AdminProductsScreen.js, mobile/src/screens/AdminSalesScreen.js - mobile admin modules',
        'mobile/src/styles/authStyles.js and mobile/src/styles/screenShell.js - shared mobile visual system',
        'mobile/scripts/check-paper-actions.js - mobile UI guard script',
        'backend/server.js - backend API startup and middleware',
        'backend/routes/authRoutes.js, backend/routes/productRoutes.js, backend/routes/saleRoutes.js, backend/routes/reportRoutes.js, backend/routes/referenceRoutes.js - backend route modules',
        'backend/middleware/authMiddleware.js - role-protection middleware',
        'backend/models/saleModel.js - sale validation, receipt generation, and voiding logic',
        'backend/config/database.js - schema initialization, seed defaults, and backups'
    ])
    add_heading(doc, 'References', level=1)
    for ref in [
        'Expo. (n.d.). Expo documentation. Retrieved March 16, 2026, from https://docs.expo.dev/',
        'Express.js. (n.d.). Express routing. Retrieved March 16, 2026, from https://expressjs.com/en/guide/routing.html',
        'Fielding, R. T. (2000). Architectural styles and the design of network-based software architectures (Doctoral dissertation, University of California, Irvine). https://roy.gbiv.com/pubs/dissertation/top.htm',
        'Jones, M., Bradley, J., & Sakimura, N. (2015, May). RFC 7519: JSON Web Token (JWT). RFC Editor. https://www.rfc-editor.org/rfc/rfc7519',
        'React. (n.d.). React. Retrieved March 16, 2026, from https://react.dev/',
        'React Navigation. (n.d.). NavigationContainer. Retrieved March 16, 2026, from https://reactnavigation.org/docs/navigation-container/',
        'React Native. (2024, October 15). Architecture overview. https://reactnative.dev/architecture/overview',
        'SQLite. (2025, November 13). Features of SQLite. https://www.sqlite.org/features.html',
        'Vite. (n.d.). Getting started. Retrieved March 16, 2026, from https://vite.dev/guide/'
    ]:
        add_text(doc, ref)

    doc.save(str(DOC_PATH))
    print(f'Generated {DOC_PATH}')


if __name__ == '__main__':
    build_document()
